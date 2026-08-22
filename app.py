import streamlit as st
import cv2
import numpy as np
from pyzbar.pyzbar import decode
from docx import Document
import io

st.set_page_config(page_title="Nhập CCCD Tự Động", page_icon="📜", layout="centered")

st.title("📜 Quét CCCD & Xuất Sơ Yếu Lý Lịch")
st.write("Dùng camera chụp trực tiếp hoặc tải lên ảnh mặt trước CCCD có mã QR.")

# Khởi tạo dữ liệu trống ban đầu
extracted_data = {
    'SO_CCCD': '',
    'HO_TEN': '',
    'NGAY_SINH': '',
    'GIOI_TINH': '',
    'DIA_CHI': '',
    'NGAY_CAP': ''
}

# --- PHẦN 1: LẤY ĐẦU VÀO ẢNH ---
# Cho phép người dùng chọn 1 trong 2 cách
input_method = st.radio("Chọn cách nhập ảnh:", ("Chụp ảnh trực tiếp", "Tải ảnh từ máy"))

final_img_bytes = None

if input_method == "Chụp ảnh trực tiếp":
    # Cách cũ: Dùng Camera input
    img_camera = st.camera_input("Chụp ảnh mặt trước CCCD")
    if img_camera is not None:
        final_img_bytes = img_camera.getvalue()
else:
    # Cách mới: Cho phép tải file ảnh lên (chọn từ thư viện ảnh điện thoại)
    img_upload = st.file_uploader("Chọn ảnh CCCD từ thiết bị", type=['jpg', 'jpeg', 'png'])
    if img_upload is not None:
        final_img_bytes = img_upload.getvalue()


# --- PHẦN 2: XỬ LÝ ẢNH & GIẢI MÃ QR ---
# Chỉ xử lý khi đã có dữ liệu ảnh (từ 1 trong 2 nguồn)
if final_img_bytes is not None:
    # Chuyển đổi byte ảnh sang định dạng OpenCV để xử lý
    cv2_img = cv2.imdecode(np.frombuffer(final_img_bytes, np.uint8), cv2.IMREAD_COLOR)

    # Giải mã QR
    # Thử giải mã trên ảnh gốc
    barcodes = decode(cv2_img)
    
    # Nếu không tìm thấy mã, thử chuyển sang ảnh xám để tăng khả năng nhận diện
    if not barcodes:
        gray = cv2.cvtColor(cv2_img, cv2.COLOR_BGR2GRAY)
        barcodes = decode(gray)

    # Nếu tìm thấy mã QR
    if barcodes:
        try:
            # Giải mã chuẩn tiếng Việt (UTF-8) từ mã QR CCCD Việt Nam
            # Cách thức: Giải mã bằng latin1 -> encode ngược lại latin1 để lấy byte gốc -> giải mã bằng utf-8
            raw_bytes = barcodes[0].data
            data_string = raw_bytes.decode('latin1').encode('latin1').decode('utf-8', errors='ignore')

            # Tách chuỗi dữ liệu bằng ký tự gạch đứng '|'
            parts = data_string.split('|')

            # Trích xuất các trường thông tin cơ bản
            if len(parts) >= 6:
                extracted_data['SO_CCCD'] = parts[0].strip()
                extracted_data['HO_TEN'] = parts[2].strip()

                # Định dạng lại ngày sinh (DDMMYYYY -> DD/MM/YYYY)
                dob_raw = parts[3].strip()
                if len(dob_raw) == 8:
                    extracted_data['NGAY_SINH'] = f"{dob_raw[:2]}/{dob_raw[2:4]}/{dob_raw[4:]}"
                else:
                    extracted_data['NGAY_SINH'] = dob_raw

                extracted_data['GIOI_TINH'] = parts[4].strip()
                extracted_data['DIA_CHI'] = parts[5].strip()

            # Trích xuất ngày cấp (phần tử thứ 7, nếu có)
            if len(parts) >= 7:
                issue_raw = parts[6].strip()
                if len(issue_raw) == 8:
                    extracted_data['NGAY_CAP'] = f"{issue_raw[:2]}/{issue_raw[2:4]}/{issue_raw[4:]}"
                else:
                    extracted_data['NGAY_CAP'] = issue_raw

            st.success("🎉 Đã quét và bóc tách thành công thông tin từ mã QR!")
        except Exception as e:
            st.error("Không thể đọc hoặc giải mã đúng định dạng mã QR CCCD.")
    else:
        st.warning("Chưa tìm thấy mã QR trong ảnh. Vui lòng đảm bảo mã QR rõ nét và không bị che khuất.")


# --- PHẦN 3: HIỂN THỊ VÀ CHỈNH SỬA THÔNG TIN ---
st.subheader("📌 Kiểm tra & Bổ sung thông tin")

col1, col2 = st.columns(2)
with col1:
    ho_ten = st.text_input("Họ và tên", value=extracted_data['HO_TEN'])
    so_cccd = st.text_input("Số CCCD", value=extracted_data['SO_CCCD'])
    ngay_sinh = st.text_input("Ngày sinh (DD/MM/YYYY)", value=extracted_data['NGAY_SINH'])
    # Xử lý giới tính: Nếu QR đọc được thì chọn đúng, nếu không thì mặc định là Nam
    gioi_tinh_index = 0 if extracted_data['GIOI_TINH'] == "Nam" else 1
    gioi_tinh = st.selectbox("Giới tính", ["Nam", "Nữ"], index=gioi_tinh_index)

with col2:
    dia_chi = st.text_input("Nơi thường trú / Địa chỉ", value=extracted_data['DIA_CHI'])
    ngay_cap = st.text_input("Ngày cấp CCCD", value=extracted_data['NGAY_CAP'])


# --- PHẦN 4: XUẤT FILE WORD ---
if st.button("🚀 Xuất file Sơ Yếu Lý Lịch (.docx)"):
    # Kiểm tra xem ít nhất đã nhập tên chưa
    if not ho_ten:
        st.error("Vui lòng nhập Họ và tên trước khi xuất file.")
    else:
        try:
            # Mở file mẫu .docx đã có sẵn trong kho lưu trữ
            doc = Document("mau_so_yeu_ly_lich.docx")

            # Tạo từ điển các từ khóa cần thay thế
            # Phải khớp chính xác với các placeholder dạng {{ TU_KHOA }} trong file mẫu
            replacements = {
                "{{ HO_TEN }}": ho_ten,
                "{{ SO_CCCD }}": so_cccd,
                "{{ NGAY_SINH }}": ngay_sinh,
                "{{ GIOI_TINH }}": gioi_tinh,
                "{{ DIA_CHI }}": dia_chi,
                "{{ NGAY_CAP }}": ngay_cap,
            }

            # Thay thế trong các đoạn văn (paragraphs)
            for p in doc.paragraphs:
                for key, val in replacements.items():
                    if key in p.text:
                        p.text = p.text.replace(key, val)

            # Thay thế trong các bảng (tables)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for key, val in replacements.items():
                                if key in p.text:
                                    p.text = p.text.replace(key, val)

            # Lưu file đã sửa vào bộ nhớ tạm (BytesIO)
            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0) # Quay về đầu file

            # Tạo nút tải file cho người dùng
            st.download_button(
                label="📥 Tải file Word về máy",
                data=bio,
                file_name=f"So_Yeu_Ly_Lich_{ho_ten.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except FileNotFoundError:
            st.error("Không tìm thấy file mẫu 'mau_so_yeu_ly_lich.docx' trong kho lưu trữ.")
        except Exception as e:
            st.error(f"Lỗi không xác định khi xuất file: {e}")
